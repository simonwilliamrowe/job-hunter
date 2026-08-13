"""
Job Hunter - Generador de paquetes por oferta usando PERSONAS (CVs
especializados) y la base de respuestas del candidato.

Regla de oro: solo se reformula/reordena lo que ya está en el perfil.
Nunca se inventan datos. La experiencia independiente crypto se presenta
SIEMPRE separada de la experiencia profesional (credibilidad, no inflación).
"""
import os
import re

from docx import Document
from docx.shared import Pt, Inches
from profile_store import load_personas

import ai


def _slug(s):
    return re.sub(r"[^a-z0-9]+", "-", (s or "").lower()).strip("-")[:40] or "oferta"


def _load_guidelines():
    """Lee data/cover_letter_guidelines.md si existe. Devuelve string o ''."""
    paths = [
        os.path.join(os.path.dirname(__file__), "data", "cover_letter_guidelines.md"),
        "data/cover_letter_guidelines.md",
    ]
    for p in paths:
        if os.path.exists(p):
            try:
                with open(p, "r", encoding="utf-8") as f:
                    return f.read()
            except Exception:
                pass
    return ""


def _guidelines_section(text, header_prefix):
    """Extrae una sección (por ## o ### header) del guidelines md."""
    if not text:
        return ""
    lines = text.splitlines()
    out, capturing = [], False
    for ln in lines:
        s = ln.strip()
        if not capturing and (s.startswith("## ") or s.startswith("### ")) and header_prefix.lower() in s.lower():
            capturing = True
            continue
        if capturing and (s.startswith("## ") or s.startswith("### ")):
            break
        if capturing:
            out.append(ln)
    return "\n".join(out).strip()


def _exp_map(profile):
    m = {}
    for e in profile.get("experience", []) + profile.get("independent_experience", []):
        m[e.get("id")] = e
    return m


def _ordered_bullets(entry, text_low):
    bullets = list(entry.get("bullets", []))
    def rank(b):
        bl = b.lower()
        return 0 if any(m in bl for m in _matched_words(text_low)) else 1
    return sorted(bullets, key=rank)


def _matched_words(text_low):
    # palabras significativas de la oferta (para ordenar bullets)
    from matcher import CRYPTO_KW, SUPPORT_KW, COMMUNITY_KW, RESEARCH_KW, FINTECH_KW, OPS_KW, TOOLS_KW
    allk = set(CRYPTO_KW + SUPPORT_KW + COMMUNITY_KW + RESEARCH_KW + FINTECH_KW + OPS_KW + TOOLS_KW)
    return [k for k in allk if re.search(r"(?<![a-z0-9])" + re.escape(k) + r"(?![a-z0-9])", text_low)][:12]


def _build_cv_text(profile, persona, job):
    lines = []
    lines.append(profile.get("name", ""))
    lines.append(persona.get("headline", profile.get("headline", "")))
    contact = " | ".join(x for x in [profile.get("email", ""), profile.get("phone", ""),
                                     profile.get("location", ""), profile.get("links", "")] if x)
    lines.append(contact)
    lines.append("")

    lines.append("SUMMARY")
    lines.append(persona.get("summary", profile.get("summary", "")))
    lines.append("")

    lines.append("SKILLS")
    lines.append(", ".join(persona.get("skills_order", [])))
    lines.append("")

    exp_map = _exp_map(profile)
    order = persona.get("experience_order", [])
    ordered = [e for e in order if e in exp_map]
    remaining = [eid for eid, e in exp_map.items() if eid not in order]
    ordered += remaining

    text_low = " ".join([job.get("title", ""), " ".join(job.get("tags", [])),
                         (job.get("description") or "")[:2500]]).lower()

    # agrupar por sección: todos los de una misma sección juntos,
    # en el orden de primera aparición que define la persona
    sections_in_order = []
    for eid in ordered:
        e = exp_map[eid]
        if not e:
            continue
        sec = "independent" if e.get("section") == "independent" else "professional"
        if sec not in sections_in_order:
            sections_in_order.append(sec)
    groups = []
    for sec in sections_in_order:
        entries = [exp_map[eid] for eid in ordered
                   if exp_map.get(eid) and
                   ("independent" if exp_map[eid].get("section") == "independent" else "professional") == sec]
        groups.append({"sec": sec, "entries": entries})

    for group in groups:
        lines.append("INDEPENDENT CRYPTO / WEB3 EXPERIENCE" if group["sec"] == "independent"
                     else "PROFESSIONAL EXPERIENCE")
        for e in group["entries"]:
            hdr = e.get("role", "")
            if e.get("company"):
                hdr = f"{hdr} — {e['company']}"
            dates = e.get("dates", "")
            if dates:
                # si la fecha ya trae paréntesis, no envolver otra vez
                hdr = f"{hdr} · {dates}" if "(" in dates else f"{hdr} ({dates})"
            lines.append(hdr)
            for b in _ordered_bullets(e, text_low):
                lines.append(f"• {b}")
            lines.append("")

    if profile.get("education"):
        lines.append("EDUCATION")
        for ed in profile["education"]:
            hdr = ed.get("degree", "")
            if ed.get("school"):
                hdr = f"{hdr} — {ed['school']}"
            dates = ed.get("dates", "")
            if dates:
                hdr = f"{hdr} · {dates}" if "(" in dates else f"{hdr} ({dates})"
            lines.append(hdr)
        lines.append("")

    if profile.get("certifications"):
        lines.append("CERTIFICATIONS")
        for c in profile["certifications"]:
            lines.append(f"• {c}")
        lines.append("")

    if profile.get("languages"):
        lines.append("LANGUAGES")
        for lang in profile["languages"]:
            l = lang.get("lang", "")
            if lang.get("level"):
                l = f"{l} — {lang['level']}"
            lines.append(l)
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def _build_docx(cv_text, path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)

    name_done = False
    section = None
    for line in cv_text.splitlines():
        line = line.rstrip()
        if not line:
            if section:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                section = None
            continue
        if not name_done:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(20)
            name_done = True
            continue
        if line.isupper() and len(line) < 40:
            section = line
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            continue
        p = doc.add_paragraph()
        if line.startswith("• "):
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(line)
        else:
            run = p.add_run(line)
            if line and "—" in line:
                run.bold = True
    doc.save(path)


def _wrap_paragraphs(text, width=95):
    """Word-wrap cada bloque separado por doble \\n a 'width' chars por línea.
    Devuelve texto con saltos de línea simples dentro de cada párrafo y
    exactamente UNA línea en blanco entre párrafos (formato carta profesional).

    Estructura esperada: header (2 líneas) \\n\\n saludo \\n\\n body... \\n\\n footer.
    El header y el footer se preservan tal cual; el saludo y los párrafos del
    body se wrappean a 'width' chars/line.
    """
    if not text:
        return ""
    import textwrap
    blocks = text.split("\n\n")
    out_blocks = []
    for i, block in enumerate(blocks):
        b = block.strip()
        if not b:
            continue
        # Primer bloque = header (nombre + contacto en 2 líneas) — preservar
        if i == 0 and "\n" in b:
            out_blocks.append(b)
        # Último bloque = footer ("Sincerely,\\nnombre") — preservar
        elif i == len(blocks) - 1 and "\n" in b:
            out_blocks.append(b)
        # Saludo: 1 línea corta, no wrappear
        elif b.startswith("Dear ") and len(b) <= width:
            out_blocks.append(b)
        # Párrafo del body: wrap
        else:
            wrapped = "\n".join(textwrap.wrap(b, width=width, break_long_words=False,
                                              break_on_hyphens=False))
            out_blocks.append(wrapped)
    return "\n\n".join(out_blocks)


def _cover_letter(profile, persona, job, ai_text=None, track_id=None):
    if ai_text:
        # Even with AI text, normalize paragraph spacing.
        return _wrap_paragraphs(ai_text, width=95)
    name = profile.get("name", "")
    contact = " | ".join(x for x in [profile.get("email", ""), profile.get("location", "")] if x)
    title = job.get("title", "")
    company = job.get("company", "")

    # Detectar si el track es trading para usar guidelines de trading
    is_trading = (track_id == "trading_entry") or ("trader" in title.lower() or "trading" in title.lower())

    # Cargar guidelines y extraer secciones relevantes
    guidelines = _load_guidelines()
    if is_trading:
        trading_section = _guidelines_section(guidelines, "2.1 Crypto Trading")
        prohibited = _guidelines_section(guidelines, "Frases prohibidas")
        connections = _guidelines_section(guidelines, "Conexiones narrativas")
    else:
        trading_section = ""
        prohibited = ""
        connections = ""

    # Construir carta según el track
    if is_trading and trading_section:
        # Carta específica para trading entry-level
        # Respetar: "knowledgeable enthusiast", "not a professional trader", WatersAbove
        watersabove = next((c for c in profile.get("certifications", [])
                           if "watersabove" in c.lower()), None)
        if watersabove:
            ws_name = watersabove.split(" — ")[0].strip()
        else:
            ws_name = "WatersAbove Crypto & Trading Course"

        # Párrafos cortos (3-5 oraciones max cada uno) con foco claro
        p1 = (
            f"I'm writing to apply for the {title} position at {company}. "
            f"This role is a natural next step for me: I've spent the last 7 "
            f"years deeply embedded in the crypto ecosystem — not as a passive "
            f"observer, but as someone who actively studies the markets, evaluates "
            f"projects through fundamentals, follows chart action, and helps others "
            f"understand the space. When I saw that you're looking for someone "
            f"with genuine interest in crypto trading, the willingness to learn, "
            f"and the discipline to grow in a structured environment, I knew this "
            f"was the right fit."
        )
        p2 = (
            f"My first professional exposure to online trading platforms was at "
            f"Binary Options Product Review (2016–2018), where I onboarded "
            f"~200–300 users to a binary options trading platform via live chat "
            f"and Skype, explaining order execution, platform mechanics and basic "
            f"risk warnings before they placed their first trades. I've since "
            f"complemented that foundation with the {ws_name}, "
            f"where I built a solid theoretical and practical base in technical "
            f"analysis, chart patterns, indicators (RSI, MACD, moving averages, "
            f"volume), Fibonacci retracements, support and resistance (micro and "
            f"macro), market and limit orders, leverage mechanics, position sizing "
            f"and risk management."
        )
        p3 = (
            f"What I bring beyond the textbook is 7+ years of continuous, "
            f"self-directed exposure to the crypto markets themselves. I "
            f"personally operate CEXs and DEXs (Binance, Kraken, KuCoin, Bitget, "
            f"SundaeSwap, Minswap, PancakeSwap). I do my own on-chain analysis via "
            f"Etherscan, Solscan and Cardanoscan — inspecting wallet activity, "
            f"transaction flows and token movements. I evaluate projects through "
            f"tokenomics, vesting schedules, market capitalisation and liquidity on "
            f"CoinGecko, CoinMarketCap and DeFiLlama, and I chart regularly on "
            f"TradingView, identifying trends, support/resistance levels and basic "
            f"chart patterns. I participated in Bitcointalk bounty programmes, "
            f"completed ~10 technical whitepaper translations EN↔ES, virtually "
            f"attended the Cardano Summit 2022, and hold the Cardano Blockchain "
            f"Fundamentals certification."
        )
        p4 = (
            f"I'm a fast learner with a process-driven mindset: at Larum Studio "
            f"(the AI & visual systems venture I co-founded in December 2025) I "
            f"built Python/OpenAI API tools to process and analyse image batches, "
            f"created an AI-assisted system for structured image hierarchy "
            f"analysis, and designed automated workflows with quality controls — "
            f"the same analytical, systematic approach that translates directly "
            f"into disciplined trading, journaling and backtesting workflows."
        )
        p5 = (
            f"I want to be transparent: I'm not a professional trader, and I don't "
            f"claim years of live P&L. What I am is a knowledgeable enthusiast "
            f"with a strong foundational understanding, 7+ years of self-directed "
            f"practice, real familiarity with the tools and platforms, and a "
            f"genuine passion for crypto markets. I'm fully remote-ready from "
            f"Paraguay (GMT-3, flexible hours, comfortable with international "
            f"teams), and the combination of comprehensive onboarding, ongoing "
            f"mentorship, hands-on platform experience and a clear development "
            f"path is exactly the structured environment I want to grow in."
        )
        p6 = (
            f"{_work_ethic_short(profile)} I'd love the chance to discuss how I "
            f"can contribute to {company}. Thank you for your time and "
            f"consideration."
        )

        # Ensamblar con doble \n entre párrafos (luego _wrap_paragraphs normaliza
        # el word-wrap a 95 chars y deja EXACTAMENTE una línea en blanco entre
        # párrafos — formato profesional de carta).
        body = "\n\n".join([p1, p2, p3, p4, p5, p6])
        # Header con nombre y contacto en líneas separadas (formato carta)
        # Usamos un separador único para que _wrap_paragraphs lo distinga
        # del body y no lo wrapee ni lo fusione con el saludo.
        header = f"{name}\n{contact}"
        footer = f"Sincerely,\n{name}"
        letter = f"{header}\n\nDear Hiring Team at {company},\n\n{body}\n\n{footer}"
        return _wrap_paragraphs(letter, width=95)

    # Carta genérica (default, todos los demás tracks)
    achievements = profile.get("achievements", [])[:3]
    ach_line = "; ".join(achievements)
    p1 = (
        f"I'm writing to apply for the {title} position at {company}. "
        f"I'm a bilingual (English/Spanish) professional based in Paraguay, fully "
        f"set up for remote work with flexible hours overlapping US/EU time zones."
    )
    p2 = (
        f"My background combines 7+ years of hands-on crypto/Web3 experience — "
        f"self-custody, wallets, exchanges, DeFi, transaction troubleshooting and "
        f"community support — with direct customer-facing work supporting 200-300 "
        f"users in financial products. Selected highlights: {ach_line}."
    )
    p3 = (
        f"{_work_ethic_short(profile)} I'd love the chance to discuss how I can "
        f"contribute to {company}. Thank you for your time and consideration."
    )
    body = "\n\n".join([p1, p2, p3])
    header = f"{name}\n{contact}"
    footer = f"Sincerely,\n{name}"
    letter = f"{header}\n\nDear Hiring Team at {company},\n\n{body}\n\n{footer}"
    return _wrap_paragraphs(letter, width=95)


def _answers_for(profile, job):
    """Base de respuestas del candidato + respuestas dinámicas de la oferta."""
    answers = dict(profile.get("answers", {}))
    parsed = None
    from matcher import parse_salary
    parsed = parse_salary(job.get("salary") or "")
    if parsed:
        answers["salary_expectation"] = (
            f"{profile.get('answers', {}).get('salary_expectation', '$1,300-$2,000/month.')} "
            f"(The listed range for this role is {parsed['raw']}.)")
    wv_long = profile.get("work_values", {}).get("long", "")
    wv_short = profile.get("work_values", {}).get("short", "")
    skills_short = ", ".join((profile.get("skills") or {}).get("support", [])[:3]) or "support and operations"
    answers.setdefault("strengths",
        f"My main strengths are my work ethic, thoroughness and bilingual communication. "
        f"I take every task seriously, no matter how small, and I hold myself to a high "
        f"standard: I don't settle for 'good enough' when I know something can be done "
        f"better. On top of that, I bring 7+ years of hands-on crypto/Web3 knowledge and "
        f"direct customer-facing experience ({skills_short}), so I can support users "
        f"clearly and calmly in both English and Spanish.")
    answers.setdefault("tell_me_about_yourself",
        f"I'm a bilingual (English/Spanish) professional based in Paraguay, fully set up "
        f"for remote work. My background combines 7+ years of hands-on crypto/Web3 "
        f"experience — wallets, exchanges, self-custody, community support — with direct "
        f"customer-facing work supporting 200-300 users in financial products. "
        f"{wv_long or ''}")

    company = job.get("company", "")
    if company:
        # Whitelist of "safe" focus areas for why_interested. We only include
        # tags that clearly map to the candidate's lanes; anything else (CSS,
        # excel, frontend, etc. from raw job tags) is dropped to avoid
        # misrepresenting the candidate's skills.
        _SAFE_FOCUS = {
            "crypto", "web3", "blockchain", "defi", "staking", "wallets",
            "customer support", "customer service", "customer success",
            "client support", "technical support", "support",
            "community", "community manager", "community support",
            "operations", "ops", "trust & safety", "trust and safety",
            "kyc", "compliance", "payments", "fintech", "banking",
            "research", "due diligence", "tokenomics",
            "ai", "content", "copy", "writing", "voice", "audio", "bilingual",
            "data entry", "admin", "virtual assistant", "va",
        }
        _NOISE_TAGS = {"all", "full time", "part time", "programming", "design",
                       "product", "sales", "marketing", "devops/sysadmin",
                       "remote", "worldwide", "anywhere", "global", "latam",
                       "english", "spanish"}
        raw_tags = [t.strip().lower() for t in (job.get("tags") or [])]
        safe_tags = [t for t in raw_tags
                     if t in _SAFE_FOCUS and t not in _NOISE_TAGS][:3]
        if safe_tags:
            focus = ", ".join(safe_tags)
        else:
            focus = "customer support, operations, and crypto/Web3 user education"
        answers["why_interested"] = (
            f"I'm interested in {company} because the role aligns with my core "
            f"strengths ({focus}) and I'm looking for a long-term remote position "
            f"where I can contribute from day one.")
    return answers


def _work_ethic_short(profile):
    wv = profile.get("work_values") or {}
    return wv.get("short", "").strip()


def _work_ethic_long(profile):
    wv = profile.get("work_values") or {}
    return wv.get("long", "").strip()


_QA_ORDER = [
    ("strengths", "What are your strengths? / Why should we hire you?"),
    ("tell_me_about_yourself", "Tell me about yourself."),
    ("years_customer_support", "How many years of customer support experience do you have?"),
    ("crypto_experience", "Do you have crypto / Web3 experience? Describe it."),
    ("work_authorization", "Are you legally authorized to work in [country]? Do you require visa sponsorship?"),
    ("salary_expectation", "What are your salary expectations? (USD/month)"),
    ("english_proficiency", "How would you rate your English proficiency?"),
    ("spanish_proficiency", "Do you speak Spanish?"),
    ("availability", "When can you start? / What is your availability?"),
    ("notice_period", "What is your notice period?"),
    ("tools_zendesk", "Have you used Zendesk / Intercom / helpdesk tools?"),
    ("remote_experience", "How much remote work experience do you have?"),
    ("education_level", "What is your highest level of education?"),
    ("why_interested", "Why do you want to work here?"),
    ("portfolio_links", "Links to portfolio / GitHub / LinkedIn"),
]


def generate_package(profile, job, track=None, out_dir="generated"):
    """Genera el paquete completo para una oferta usando la persona del track."""
    personas = load_personas()
    from matcher import score_offer
    scoring = score_offer(job, profile, personas)
    track_id = track or scoring["track"]
    persona = personas.get(track_id, personas.get(scoring["persona"], {}))
    if not persona and scoring["persona"] in personas:
        persona = personas[scoring["persona"]]

    safe_dir = os.path.join(out_dir, _slug(job["id"]) or "oferta")
    os.makedirs(safe_dir, exist_ok=True)

    cv_text = _build_cv_text(profile, persona, job)
    slug = _slug(job["company"])

    cv_docx = os.path.join(safe_dir, f"CV_{slug}.docx")
    _build_docx(cv_text, cv_docx)
    cv_txt = os.path.join(safe_dir, f"CV_{slug}.txt")
    with open(cv_txt, "w", encoding="utf-8") as f:
        f.write(cv_text)

    cv_pdf = os.path.join(safe_dir, f"CV_{slug}.pdf")
    try:
        import pdfgen
        pdfgen.build_pdf(profile, persona, cv_pdf)
    except Exception as e:  # noqa: BLE001
        print(f"[cvgen] PDF opcional no generado: {e}")
        cv_pdf = None

    ai_cover = ai.ai_rewrite(profile, job, "cover")
    cover = _cover_letter(profile, persona, job, ai_text=ai_cover, track_id=track_id)
    cover_path = os.path.join(safe_dir, "cover_letter.txt")
    with open(cover_path, "w", encoding="utf-8") as f:
        f.write(cover)

    answers = _answers_for(profile, job)
    answers_path = os.path.join(safe_dir, "application_answers.txt")
    with open(answers_path, "w", encoding="utf-8") as f:
        for key, q in _QA_ORDER:
            if key in answers and answers[key]:
                f.write(f"Q: {q}\nA: {answers[key]}\n\n")

    return {
        "offer_id": job["id"],
        "dir": _slug(job["id"]) or "oferta",
        "track": track_id,
        "track_label": scoring["track_label"],
        "persona": persona.get("label", track_id),
        "score": scoring["score"],
        "band": scoring["band"],
        "emoji": scoring["emoji"],
        "files": [f for f in [
            {"name": os.path.basename(cv_docx), "path": cv_docx, "kind": "CV (Word, ATS-friendly)"},
            {"name": os.path.basename(cv_txt), "path": cv_txt, "kind": "CV (texto, para formularios)"},
            {"name": os.path.basename(cv_pdf), "path": cv_pdf, "kind": "CV (PDF, para adjuntar)"} if cv_pdf else None,
            {"name": "cover_letter.txt", "path": cover_path, "kind": "Carta de presentación"},
            {"name": "application_answers.txt", "path": answers_path, "kind": "Respuestas del formulario"},
        ] if f],
        "answers": [{"question": q, "answer": answers.get(key, "")} for key, q in _QA_ORDER if answers.get(key)],
        "matched_skills": scoring.get("matched", []),
        "ai_tailored": bool(ai_cover),
    }
